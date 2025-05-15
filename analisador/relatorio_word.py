from docx import Document

class WordReport:
    @staticmethod
    def gerar(caminho_saida, analises):
        doc = Document()
        doc.add_heading("Análise de Artigos sobre Nanopartículas", 0)

        for entrada in analises:
            doc.add_heading(entrada["titulo"], level=1)
            tabela = doc.add_table(rows=8, cols=2)
            tabela.style = 'Table Grid'

            linhas = [
                ("Nanopartículas", entrada["dados"]["nanopartículas"]),
                ("Utilidades", entrada["dados"]["utilidades"]),
                ("Patentes", entrada["dados"]["patentes"]),
                ("Tamanho das Partículas", entrada["dados"]["tamanho"]),
                ("Técnica de Fabricação", entrada["dados"]["técnica"]),
                ("Composição / Matriz Polimérica", entrada["dados"]["composição"]),
                ("Aplicações Sugeridas", entrada["dados"]["aplicações"]),
                ("Fonte", entrada["titulo"])
            ]

            for i, (campo, valor) in enumerate(linhas):
                tabela.cell(i, 0).text = campo
                tabela.cell(i, 1).text = valor

            doc.add_paragraph("\n")

        doc.save(caminho_saida)
